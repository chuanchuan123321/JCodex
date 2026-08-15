"""Extended tool executor with document reading capabilities"""

import base64
import copy
import hashlib
import json
import os
import re
import signal
import subprocess
import time
import threading
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple

import requests
from langchain_community.agent_toolkits import FileManagementToolkit

from agent.tools.shell import ShellTool, decode_stream
from agent.tools.url_fetch import fetch_url
from agent.tools.file import FileTool
from agent.tools.time_tool import TimeTool
from agent.tools.pdf_tool import PDFTool
from agent.tools.skill_tool import SkillTool
from agent.tools.grep import execute_grep, get_grep_tool_definition
from agent.tools.edit import execute_edit, get_edit_tool_definition
from agent.tools.websearch import execute_websearch, get_websearch_tool_definition
from agent.tools.preview import PreviewManager
from agent.tools.plan import PlanTool
from agent.core.multi_agent import MULTI_AGENT_TOOL_NAMES
from agent.core.tool_result import ToolExecutionResult


MAX_TASK_IMAGE_BYTES = 12 * 1024 * 1024
_TASK_IMAGE_MIME_TYPES = {"image/png", "image/jpeg", "image/webp"}


def _vision_tools_enabled() -> bool:
    """Return whether the multimodal view_image tooling is enabled."""
    return os.getenv("MODEL_SUPPORTS_VISION", "true").strip().lower() not in {
        "0",
        "false",
        "no",
        "off",
    }


def strip_disabled_vision_prompt(prompt: str) -> str:
    """Remove the view_image tool line when multimodal vision is off."""
    if _vision_tools_enabled():
        return str(prompt or "")
    return "\n".join(
        line
        for line in str(prompt or "").splitlines()
        if not line.lstrip().startswith("- `view_image`")
    )


# A scoped child may use these file tools only when every mutation target stays
# beneath one of its coordinator-assigned roots. Each inner tuple represents
# alternative parameter spellings for one required target.
_SCOPED_MUTATION_TARGETS = {
    "write": (("path",),),
    "edit": (("filePath", "file_path", "path"),),
    "search_replace": (("filePath", "file_path", "path"),),
    "generate_pdf": (("output_path",),),
}

# Mutations that patch existing content must only run against the version the
# model actually read. Whole-file writes replace everything, so they carry no
# stale concern of their own and are excluded here.
_STALE_CHECKED_TARGETS = {
    "edit": (("filePath", "file_path", "path"),),
    "search_replace": (("filePath", "file_path", "path"),),
}

# Shell commands cannot be constrained reliably with path preflight checks: a
# command can write through redirection, child processes, scripts, or utilities.
_SCOPED_COMMAND_TOOLS = frozenset(
    {"bash", "shell", "run_terminal_cmd", "pwsh", "monitor"}
)
_MUTATION_SCOPE_UNSET = object()


def _detect_supported_image_mime(content: bytes) -> Optional[str]:
    """Identify image bytes accepted by the model-facing image viewer."""
    if content.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image/png"
    if content.startswith(b"\xff\xd8\xff"):
        return "image/jpeg"
    if (
        len(content) >= 12
        and content.startswith(b"RIFF")
        and content[8:12] == b"WEBP"
    ):
        return "image/webp"
    return None


@dataclass(frozen=True)
class _TaskImage:
    """One task-scoped image attachment available to the image viewer tool."""

    path: Path
    name: str
    mime_type: str
    size: int
    digest: str


class ExtendedToolExecutor:
    """Execute tools with extended capabilities including document reading"""

    def __init__(
        self,
        skills_loader=None,
        preview_manager: Optional[PreviewManager] = None,
        conversation_id: Optional[str] = None,
        message_id: Optional[str] = None,
        project_root: Optional[str] = None,
        workspace_root: Optional[str] = None,
        protected_root: Optional[str] = None,
        data_root: Optional[str] = None,
        restrict_reads_to_project: bool = False,
    ):
        # self.shell_tool is created below, after workspace_temp_root resolves.
        self.file_tool = FileTool()
        self.pdf_tool = PDFTool()
        self.skill_tool = SkillTool(skills_loader) if skills_loader else None
        self.project_root = Path(
            project_root or Path(__file__).resolve().parents[2]
        ).resolve()
        self.workspace_root = Path(
            workspace_root or self.project_root / "workspace"
        ).resolve()
        self.workspace_temp_root = (
            self.workspace_root / "temp"
        ).resolve()
        # Truncated shell output spills here so the model can read the tail
        # that was cut from the returned result.
        self.shell_tool = ShellTool(spill_dir=str(self.workspace_temp_root))
        # Group background-command output logs under temp/tasks.
        self.task_log_root = self.workspace_temp_root / "tasks"
        try:
            self.task_log_root.mkdir(parents=True, exist_ok=True)
            for old_log in self.workspace_temp_root.glob("task-*.log"):
                old_log.rename(self.task_log_root / old_log.name)
        except OSError:
            pass
        self.workspace_output_root = (
            self.workspace_root / "output"
        ).resolve()
        self.protected_root = Path(
            protected_root or Path(__file__).resolve().parents[2]
        ).resolve()
        # Runtime data root (user Application Support dir in the packaged app).
        # File mutations are restricted to workspace/output and workspace/temp
        # under this root, while the source/bundle tree stays protected too.
        self.data_root = Path(data_root or self.protected_root).resolve()
        self._protected_regions = tuple(
            dict.fromkeys(
                root
                for root in (self.protected_root, self.data_root)
                if root
            )
        )
        self.protected_write_roots = (
            self.workspace_output_root,
            self.workspace_temp_root,
        )
        # Kept for callers using the old constructor; local reads are unrestricted.
        self.restrict_reads_to_project = bool(restrict_reads_to_project)
        self._file_search_tool = next(
            tool
            for tool in FileManagementToolkit(
                root_dir=str(self.project_root), selected_tools=["file_search"]
            ).get_tools()
            if tool.name == "file_search"
        )
        self.preview_manager = preview_manager or PreviewManager(
            self.project_root
        )
        self.conversation_id = conversation_id
        self.message_id = message_id
        self._plan_tools: Dict[str, PlanTool] = {}
        self._plan_tools_lock = threading.RLock()
        self._todo_states: Dict[str, Dict[str, Any]] = {}
        self._todo_state_lock = threading.RLock()
        self._loaded_skills: Dict[str, set[str]] = {}
        self._loaded_skills_lock = threading.RLock()
        self._task_images: Dict[str, Dict[str, _TaskImage]] = {}
        self._task_images_lock = threading.RLock()
        self._task_reference_roots: Dict[str, tuple[Path, ...]] = {}
        self._task_reference_roots_lock = threading.RLock()
        self.memory_store = None
        # File version fingerprints (mtime_ns, size) recorded on each successful
        # read, so a later edit can detect that the target changed since it was
        # read and demand a re-read instead of patching stale content. Whole-file
        # writes replace everything and are exempt from this check.
        self._read_versions: Dict[str, Tuple[int, int]] = {}
        self._background_tasks: Dict[str, Dict[str, Any]] = {}
        self._background_tasks_lock = threading.RLock()
        self._scheduled_tasks: Dict[str, Dict[str, Any]] = {}
        self._scheduled_tasks_lock = threading.RLock()
        self.scheduled_prompt_callback: Optional[Callable[[str], Any]] = None
        self._goal_state: Dict[str, Any] = {"status": "active", "message": ""}
        self.tools: Dict[str, Callable] = {
            "bash": self.execute_shell,
            "pwsh": self.execute_shell,
            "read": self.execute_file_read,
            "glob": self.execute_glob,
            "grep": self.execute_grep,
            "edit": self.execute_edit,
            "write": self.execute_file_write,
            "websearch": execute_websearch,
            "read_url": self.execute_read_url,
            "send_file": self.execute_send_file,
            "generate_pdf": self.execute_generate_pdf,
            "load_skill": self.execute_load_skill,
            "question": self.execute_question,
            "update_plan": self.execute_update_plan,
            "view_image": self.execute_view_image,
            "project_preview": self.execute_project_preview,
            # Grok Build-compatible model-facing names.
            "run_terminal_cmd": self.execute_shell,
            "read_file": self.execute_file_read,
            "search_replace": self.execute_edit,
            "list_dir": self.execute_file_list,
            "web_fetch": self.execute_read_url,
            "ask_user_question": self.execute_question,
            "todo_write": self.execute_todo_write,
            "memory_search": self.execute_memory_search,
            "memory_get": self.execute_memory_get,
            "get_task_output": self.execute_get_task_output,
            "get_terminal_command_output": self.execute_get_task_output,
            "wait_tasks": self.execute_get_task_output,
            "kill_task": self.execute_kill_task,
            "kill_terminal_command": self.execute_kill_task,
            "monitor": self.execute_monitor,
            "search_tool": self.execute_search_tool,
            "use_tool": self.execute_use_tool,
            "wait_tasks": self.execute_get_task_output,
            "scheduler_create": self.execute_scheduler_create,
            "scheduler_delete": self.execute_scheduler_delete,
            "scheduler_list": self.execute_scheduler_list,
            "update_goal": self.execute_update_goal,
            # Defensive aliases kept out of the model tool list: shell/file_read
            # guard legacy approvals and web_search is a Grok-compatible name.
            "shell": self.execute_shell,
            "file_read": self.execute_file_read,
            "web_search": execute_websearch,
        }

    def get_available_tools(
        self, *, include_gateway_tools: bool = False
    ) -> List[Dict[str, Any]]:
        """Return model tools, exposing gateway-only actions only when requested."""
        tools = [
            {
                "type": "function",
                "function": {
                    "name": "view_image",
                    "description": "View one PNG, JPEG, or WebP image available to the current task. Use either a full path listed in the current task's image attachment manifest or a full path inside workspace/temp or workspace/output. The image is sent to the model only for this task run; arbitrary local image paths are rejected.",
                    "parameters": {
                        "type": "object",
                        "properties": {
                            "path": {
                                "type": "string",
                                "description": "Exact full path of a current-task image attachment or a supported image inside workspace/temp or workspace/output",
                            },
                        },
                        "required": ["path"],
                    },
                },
            },
            {
                "type": "function",
                "function": {
                    "name": "project_preview",
                    "description": "Start, inspect, or stop a persistent loopback-only Web project preview. After completing and verifying a website or Web app, proactively call action=start so the user receives the preview component before the task ends. Use this instead of bash for long-running development servers. For Python static sites, `python3 -m http.server` is enough: the preview manager injects the managed port and 127.0.0.1 binding automatically, and on Windows packaged builds the bundled runtime serves static sites the same way, so no python3 is required.",
                    "parameters": {
                        "type": "object",
                        "properties": {
                            "action": {
                                "type": "string",
                                "enum": ["start", "status", "stop"],
                                "description": "Preview lifecycle action",
                            },
                            "command": {
                                "type": "string",
                                "description": "Start command. Pass injected $HOST/$PORT on this command (not inside package.json scripts), e.g. npm run dev -- --host $HOST --port $PORT. Python `-m http.server` commands are normalized automatically. Required for start.",
                            },
                            "workdir": {
                                "type": "string",
                                "description": "Working directory to preview. Relative paths resolve from the active project or default task root; absolute local directories are also supported.",
                            },
                            "name": {
                                "type": "string",
                                "description": "Short project name displayed in the preview card",
                            },
                            "port": {
                                "type": "integer",
                                "minimum": 0,
                                "maximum": 65535,
                                "description": "Loopback port; use 0 or omit to allocate one automatically",
                            },
                            "health_path": {
                                "type": "string",
                                "description": "Local HTTP path used for readiness checks (default: /)",
                            },
                            "entry_path": {
                                "type": "string",
                                "description": "Same-origin page path to open in the preview, such as /jcodex.html or /nested/demo.html. Do not pass a full URL. When omitted, root index.html or a single root HTML file is selected automatically.",
                            },
                            "startup_timeout": {
                                "type": "number",
                                "minimum": 1,
                                "maximum": 120,
                                "description": "Maximum seconds to wait for the server to become reachable",
                            },
                            "preview_id": {
                                "type": "string",
                                "description": "Preview identifier returned by start; required for stop and optional for status",
                            },
                        },
                        "required": ["action"],
                    },
                },
            },
            *self._shell_tool_definitions(),
            {
                "type": "function",
                "function": {
                    "name": "read",
                    "description": "Read a file or directory from the local filesystem. PDF, Word, and Excel files are parsed automatically. Image files cannot be read; the read tool only supports document files. Reads at most 1000 lines per call. A window cut short by the output budget ends with `(Showing lines X-Y. Use offset=Y+1 to continue.)` — follow it to keep reading; lines longer than 2000 characters are truncated with an explicit suffix. For large files, use grep first, then read a focused range with offset and limit.",
                    "parameters": {
                        "type": "object",
                        "properties": {
                            "filePath": {
                                "type": "string",
                                "description": "The absolute path to the file or directory to read",
                            },
                            "offset": {
                                "type": "number",
                                "description": "The line number to start reading from (1-indexed)",
                            },
                            "limit": {
                                "type": "number",
                                "description": "The maximum number of lines to read (defaults to and is capped at 1000)",
                            },
                        },
                        "required": ["filePath"],
                    },
                },
            },
            self.get_glob_tool_definition(),
            get_grep_tool_definition(),
            get_edit_tool_definition(),
            {
                "type": "function",
                "function": {
                    "name": "write",
                    "description": "Writes a file to the local filesystem. This tool will overwrite existing files.",
                    "parameters": {
                        "type": "object",
                        "properties": {
                            "content": {
                                "type": "string",
                                "description": "The content to write to the file",
                            },
                            "path": {
                                "type": "string",
                                "description": "The absolute path to the file to write",
                            },
                        },
                        "required": ["content", "path"],
                    },
                },
            },
            {
                "type": "function",
                "function": {
                    "name": "generate_pdf",
                    "description": "Generate PDF from Markdown, text, HTML, or Word documents",
                    "parameters": {
                        "type": "object",
                        "properties": {
                            "input_path": {
                                "type": "string",
                                "description": "Input file path",
                            },
                            "output_path": {
                                "type": "string",
                                "description": "Output PDF file path",
                            },
                            "format": {
                                "type": "string",
                                "description": "Input format (markdown/text/html/docx)",
                            },
                        },
                        "required": ["input_path", "output_path"],
                    },
                },
            },
            {
                "type": "function",
                "function": {
                    "name": "load_skill",
                    "description": "Load a skill's complete content to get detailed guidance and instructions",
                    "parameters": {
                        "type": "object",
                        "properties": {
                            "skill_name": {
                                "type": "string",
                                "description": "Name of the skill to load",
                            },
                        },
                        "required": ["skill_name"],
                    },
                },
            },
            {
                "type": "function",
                "function": {
                    "name": "question",
                    "description": "Pause execution and ask the user one or more selectable questions. Do not continue until the user submits answers. Set multiple=true for every multi-select question. Set allow_free_text=true when the user may add text; never describe a capability in the question text without setting its matching field.",
                    "parameters": {
                        "type": "object",
                        "properties": {
                            "questions": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "header": {
                                            "type": "string",
                                            "description": "Very short label (max 30 chars)",
                                        },
                                        "multiple": {
                                            "type": "boolean",
                                            "description": "Required. True only when the user may select more than one option; false for exactly one choice.",
                                        },
                                        "selection_required": {
                                            "type": "boolean",
                                            "description": "Whether a listed option must be selected. Defaults to true; when allow_free_text is true, typed text can also satisfy the question.",
                                        },
                                        "allow_free_text": {
                                            "type": "boolean",
                                            "description": "Show a text supplement field below this question.",
                                        },
                                        "free_text_label": {
                                            "type": "string",
                                            "description": "Short label for the optional text supplement field.",
                                        },
                                        "free_text_placeholder": {
                                            "type": "string",
                                            "description": "Helpful example or placeholder for the text supplement field.",
                                        },
                                        "free_text_required": {
                                            "type": "boolean",
                                            "description": "Require text in the supplement field before submission.",
                                        },
                                        "options": {
                                            "type": "array",
                                            "items": {
                                                "type": "object",
                                                "properties": {
                                                    "description": {
                                                        "type": "string",
                                                        "description": "Explanation of choice",
                                                    },
                                                    "label": {
                                                        "type": "string",
                                                        "description": "Display text",
                                                    },
                                                },
                                                "required": ["label", "description"],
                                            },
                                        },
                                        "question": {
                                            "type": "string",
                                            "description": "Complete question",
                                        },
                                    },
                                    "required": ["question", "header", "multiple", "options"],
                                },
                            },
                        },
                        "required": ["questions"],
                    },
                },
            },
        ]

        if include_gateway_tools:
            tools.append(
                {
                    "type": "function",
                    "function": {
                        "name": "send_file",
                        "description": "Send a file to the user through the active gateway channel.",
                        "parameters": {
                            "type": "object",
                            "properties": {
                                "path": {
                                    "type": "string",
                                    "description": "Path to the file to send",
                                },
                            },
                            "required": ["path"],
                        },
                    },
                }
            )
        tools.extend(self._grok_compatible_tool_definitions())
        tools.extend(self._multi_agent_tool_definitions())
        if not _vision_tools_enabled():
            tools = [
                tool
                for tool in tools
                if tool.get("function", {}).get("name") != "view_image"
            ]
        return tools

    def _shell_tool_definitions(self) -> List[Dict[str, Any]]:
        """Return the model-facing shell tool for this platform.

        Platform split: macOS/Linux expose the ``bash`` tool with a
        POSIX-dialect contract; Windows exposes the ``pwsh`` tool with a
        PowerShell-dialect contract (when PowerShell is available) so the model
        never receives commands in the wrong dialect. The legacy ``bash``/
        ``run_terminal_cmd`` names stay registered for dispatch so older
        checkpoints can still resume, but they are not advertised to the model
        on Windows.
        """
        if os.name == "nt" and ShellTool._resolve_pwsh() is not None:
            return [self._pwsh_tool_definition()]
        return [self._bash_tool_definition()]

    @staticmethod
    def _shell_tool_parameters() -> Dict[str, Any]:
        """Shared parameter schema for the bash/pwsh tools."""
        return {
            "type": "object",
            "properties": {
                "command": {
                    "type": "string",
                    "description": "The command to execute",
                },
                "timeout": {
                    "type": "number",
                    "description": "Optional timeout in milliseconds",
                },
                "workdir": {
                    "type": "string",
                    "description": "The working directory to run the command in. Defaults to the session workspace; a relative path is resolved against it.",
                },
                "description": {
                    "type": "string",
                    "description": "Clear, concise description of what this command does in active voice, 5-10 words (shown in the UI). Examples: \"ls\" → \"List files in current directory\"; \"git status\" → \"Show working tree status\"; \"npm install\" → \"Install package dependencies\".",
                },
                "is_background": {
                    "type": "boolean",
                    "description": "Run a long-lived command in the background and return a task ID; collect output with get_task_output and stop it with kill_task",
                },
            },
            "required": ["command"],
        }

    def _bash_tool_definition(self) -> Dict[str, Any]:
        """The POSIX shell tool contract (macOS/Linux)."""
        return {
            "type": "function",
            "function": {
                "name": "bash",
                "description": (
                    "Execute a command in the system shell (bash on macOS, sh on Linux) and return "
                    "its stdout/stderr. Each call runs in a fresh shell: no state (cwd, variables, "
                    "functions) persists between calls — pass `workdir` instead of using `cd`. "
                    "Non-zero exits are reported as `[exit code: N]`: a non-zero exit is a report, "
                    "not an error (grep with no matches, diff with differences, and curl -f exit "
                    "non-zero normally) — inspect the output and decide how to react. Timeouts, "
                    "cancellations, and signal kills appear as `[timed out after Ns]`, `[cancelled]`, "
                    "and `[killed by signal: X]`. Truncated output appends "
                    "`[output truncated; full output: <path>]`; read the full log file when the tail "
                    "is not enough. Check the `[exit code: N]` marker on every result; investigate "
                    "failures before moving on."
                ),
                "parameters": self._shell_tool_parameters(),
            },
        }

    def _pwsh_tool_definition(self) -> Dict[str, Any]:
        """The Windows PowerShell tool contract."""
        return {
            "type": "function",
            "function": {
                "name": "pwsh",
                "description": (
                    "Execute a PowerShell command via `pwsh -NoProfile -NonInteractive -Command` "
                    "(falling back to `powershell` when pwsh is absent) and return its stdout/stderr. "
                    "Each call runs in a fresh PowerShell process: no state (cwd, variables, functions) "
                    "persists between calls — pass `workdir` instead of using `cd`. Paths use native "
                    "Windows form (`C:\\...`); read environment variables with `$env:NAME`; use "
                    "PowerShell cmdlets (Get-ChildItem, Get-Content, Copy-Item, Move-Item, "
                    "Remove-Item, Select-String, Where-Object) rather than POSIX commands. "
                    "Non-zero exits are reported as `[exit code: N]`: a non-zero exit is a report, not "
                    "an error — inspect the output and decide how to react. On Windows a force-killed "
                    "command settles as `[exit code: 1]` without a signal marker — treat it as an "
                    "interruption, not a command failure. Timeouts appear as `[timed out after Ns]`; "
                    "truncated output appends `[output truncated; full output: <path>]`. Check the "
                    "`[exit code: N]` marker on every result; investigate failures before moving on."
                ),
                "parameters": self._shell_tool_parameters(),
            },
        }

    def _grok_compatible_tool_definitions(self) -> List[Dict[str, Any]]:
        """Return executable Grok Build-compatible aliases and extensions."""
        definitions = []
        aliases = {
            "web_search": (
                "websearch",
                "Search the public web for current or unknown information. Returns an optional summary answer and a list of source URLs with snippets. Set include_images to true only for visual queries (images, photos, products); images come back as related links with little or no description. Use focused queries and avoid repeating equivalent searches.",
            ),
            "web_fetch": (
                "read_url",
                "Fetches content from a specific HTTP(S) URL and returns it decoded to text. Only same-host redirects are followed automatically; a redirect to a different host is refused — fetch that URL directly. Non-text responses (images, archives, downloads) are rejected. Long pages are truncated with a notice at the end; fetch a more specific URL or section for the full text.",
            ),
        }
        by_name = {
            tool["function"]["name"]: tool
            for tool in definitions
            + [get_edit_tool_definition(), get_websearch_tool_definition()]
        }
        base_tools = {tool["function"]["name"]: tool for tool in self._legacy_public_tools()}
        base_tools.update(by_name)
        for name, (source, description) in aliases.items():
            definition = copy.deepcopy(base_tools[source])
            definition["function"]["name"] = name
            definition["function"]["description"] = description
            definitions.append(definition)

        definitions.extend(
            [
                self._list_dir_definition(),
                self._todo_write_definition(),
                self._memory_search_definition(),
                self._memory_get_definition(),
                self._task_output_definition(),
                self._wait_tasks_definition(),
                self._kill_task_definition(),
                self._monitor_definition(),
                self._search_tool_definition(),
                self._use_tool_definition(),
                self._scheduler_create_definition(),
                self._scheduler_delete_definition(),
                self._scheduler_list_definition(),
                self._update_goal_definition(),
            ]
        )
        existing = {tool["function"]["name"] for tool in self._legacy_public_tools()}
        return [
            tool for tool in definitions if tool["function"]["name"] not in existing
        ]

    def _legacy_public_tools(self) -> List[Dict[str, Any]]:
        """Build only definitions needed as sources for Grok-compatible aliases."""
        return [
            {
                "type": "function",
                "function": {
                    "name": "bash",
                    "description": "Run a terminal command.",
                    "parameters": {
                        "type": "object",
                        "properties": {
                            "command": {"type": "string"},
                            "timeout": {"type": "number"},
                            "workdir": {"type": "string"},
                            "description": {"type": "string"},
                            "is_background": {"type": "boolean"},
                        },
                        "required": ["command"],
                    },
                },
            },
            {
                "type": "function",
                "function": {
                    "name": "read",
                    "description": "Read a file.",
                    "parameters": {
                        "type": "object",
                        "properties": {
                            "path": {"type": "string"},
                            "from": {"type": "integer", "minimum": 0},
                            "lines": {"type": "integer", "minimum": 1},
                        },
                        "required": ["path"],
                    },
                },
            },
            {
                "type": "function",
                "function": {
                    "name": "read_url",
                    "description": "Fetch a URL.",
                    "parameters": {
                        "type": "object",
                        "properties": {"url": {"type": "string"}},
                        "required": ["url"],
                    },
                },
            },
            {
                "type": "function",
                "function": {
                    "name": "question",
                    "description": "Ask the user questions.",
                    "parameters": {
                        "type": "object",
                        "properties": {"questions": {"type": "array"}},
                        "required": ["questions"],
                    },
                },
            },
        ]

    @staticmethod
    def _list_dir_definition() -> Dict[str, Any]:
        return {
            "type": "function",
            "function": {
                "name": "list_dir",
                "description": "List directory contents relative to the project root or by absolute path.",
                "parameters": {
                    "type": "object",
                    "properties": {"target_directory": {"type": "string"}},
                    "required": ["target_directory"],
                },
            },
        }

    @staticmethod
    def _todo_write_definition() -> Dict[str, Any]:
        return {
            "type": "function",
            "function": {
                "name": "todo_write",
                "description": "Create and manage a structured task list. Use for tasks with three or more steps.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "merge": {"type": "boolean", "default": True},
                        "todos": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "id": {"type": "string"},
                                    "content": {"type": "string"},
                                    "status": {
                                        "type": "string",
                                        "enum": ["pending", "in_progress", "completed", "cancelled"],
                                    },
                                },
                                "required": ["id"],
                            },
                        },
                    },
                    "required": ["todos"],
                },
            },
        }

    @staticmethod
    def _memory_search_definition() -> Dict[str, Any]:
        return {
            "type": "function",
            "function": {
                "name": "memory_search",
                "description": "Search cross-session memory for ranked global, workspace, and session knowledge chunks.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "query": {"type": "string"},
                        "max_results": {"type": "integer", "minimum": 1},
                        "min_score": {"type": "number", "minimum": 0, "maximum": 1},
                    },
                    "required": ["query"],
                },
            },
        }

    @staticmethod
    def _memory_get_definition() -> Dict[str, Any]:
        return {
            "type": "function",
            "function": {
                "name": "memory_get",
                "description": "Read a memory Markdown file returned by memory_search with optional line bounds.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string"},
                        "from": {"type": "integer", "minimum": 0},
                        "lines": {"type": "integer", "minimum": 1},
                    },
                    "required": ["path"],
                },
            },
        }

    @staticmethod
    def _task_output_definition() -> Dict[str, Any]:
        return {
            "type": "function",
            "function": {
                "name": "get_task_output",
                "description": "Get output and status from one or more background terminal commands.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "task_ids": {"type": "array", "items": {"type": "string"}},
                        "timeout_ms": {"type": "integer", "minimum": 0, "maximum": 600000},
                    },
                    "required": ["task_ids"],
                },
            },
        }

    @staticmethod
    def _wait_tasks_definition() -> Dict[str, Any]:
        definition = copy.deepcopy(ExtendedToolExecutor._task_output_definition())
        definition["function"]["name"] = "wait_tasks"
        definition["function"]["description"] = (
            "Compatibility wait tool for background task IDs. Prefer get_task_output."
        )
        return definition

    @staticmethod
    def _kill_task_definition() -> Dict[str, Any]:
        return {
            "type": "function",
            "function": {
                "name": "kill_task",
                "description": "Terminate a background terminal command by task ID.",
                "parameters": {
                    "type": "object",
                    "properties": {"task_id": {"type": "string"}},
                    "required": ["task_id"],
                },
            },
        }

    @staticmethod
    def _monitor_definition() -> Dict[str, Any]:
        return {
            "type": "function",
            "function": {
                "name": "monitor",
                "description": "Start a long-running command in the background and return a task ID for monitoring.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "command": {"type": "string"},
                        "workdir": {"type": "string"},
                        "description": {"type": "string"},
                    },
                    "required": ["command"],
                },
            },
        }

    @staticmethod
    def _search_tool_definition() -> Dict[str, Any]:
        return {
            "type": "function",
            "function": {
                "name": "search_tool",
                "description": "Search the active tool catalog by capability or name.",
                "parameters": {
                    "type": "object",
                    "properties": {"query": {"type": "string"}},
                    "required": ["query"],
                },
            },
        }

    @staticmethod
    def _use_tool_definition() -> Dict[str, Any]:
        return {
            "type": "function",
            "function": {
                "name": "use_tool",
                "description": "Invoke a tool discovered with search_tool using its exact input schema.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "tool_name": {"type": "string"},
                        "tool_input": {"type": "object"},
                    },
                    "required": ["tool_name", "tool_input"],
                },
            },
        }

    @staticmethod
    def _scheduler_create_definition() -> Dict[str, Any]:
        return {
            "type": "function",
            "function": {
                "name": "scheduler_create",
                "description": "Create or update a recurring scheduled prompt. Intervals use 60s, 5m, 2h, or 1d syntax. When it fires, the prompt is run as a new task inside the current conversation (the AI continues working there); the user can view and delete it from the settings sidebar.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "task_id": {"type": "string"},
                        "interval": {"type": "string"},
                        "prompt": {"type": "string"},
                        "recurring": {"type": "boolean", "default": True},
                        "durable": {"type": "boolean", "default": False},
                        "foreground": {"type": "boolean", "default": False},
                        "fire_immediately": {"type": "boolean", "default": False},
                    },
                },
            },
        }

    @staticmethod
    def _scheduler_delete_definition() -> Dict[str, Any]:
        return {
            "type": "function",
            "function": {
                "name": "scheduler_delete",
                "description": "Cancel a scheduled task by ID.",
                "parameters": {
                    "type": "object",
                    "properties": {"task_id": {"type": "string"}},
                    "required": ["task_id"],
                },
            },
        }

    @staticmethod
    def _scheduler_list_definition() -> Dict[str, Any]:
        return {
            "type": "function",
            "function": {
                "name": "scheduler_list",
                "description": "List active scheduled tasks and their next fire times.",
                "parameters": {"type": "object", "properties": {}},
            },
        }

    @staticmethod
    def _update_goal_definition() -> Dict[str, Any]:
        return {
            "type": "function",
            "function": {
                "name": "update_goal",
                "description": "Report goal progress, completion, or a genuine blocker.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "completed": {"type": "boolean"},
                        "message": {"type": "string"},
                        "blocked_reason": {"type": "string"},
                    },
                },
            },
        }

    @staticmethod
    def _multi_agent_tool_definitions() -> List[Dict[str, Any]]:
        """Return tools dispatched by the active desktop collaboration runtime."""
        return [
            {
                "type": "function",
                "function": {
                    "name": "spawn_agent",
                    "description": (
                        "Create one isolated child agent for a concrete independent "
                        "subtask. Give it a short visible name and role. The child "
                        "receives only the explicit task and context supplied here."
                    ),
                    "parameters": {
                        "type": "object",
                        "properties": {
                            "name": {
                                "type": "string",
                                "description": "Short name displayed in the collaboration UI",
                            },
                            "role": {
                                "type": "string",
                                "description": "The child's specialization and responsibility",
                            },
                            "task": {
                                "type": "string",
                                "description": "One bounded deliverable for this child",
                            },
                            "context": {
                                "type": "string",
                                "description": "Only the background and constraints this child needs",
                            },
                            "write_access": {
                                "type": "boolean",
                                "description": (
                                    "Set true for an implementation child that must "
                                    "create or modify files. Keep false for research "
                                    "or review-only work. Requires non-overlapping "
                                    "write_paths."
                                ),
                            },
                            "write_paths": {
                                "type": "array",
                                "items": {"type": "string"},
                                "description": (
                                    "Required when write_access is true. Use explicit "
                                    "project-relative files or directories owned only by "
                                    "this child, such as src/api/ or tests/. The child "
                                    "may modify only these paths. Directory ownership "
                                    "is recursive; do not append glob markers such as "
                                    "* or **."
                                ),
                            },
                            "workdir": {
                                "type": "string",
                                "description": (
                                    "Working directory used to resolve relative "
                                    "write_paths and child file-tool paths. Set this "
                                    "for generated projects outside the active project, "
                                    "for example workspace/output/my-app. The "
                                    "coordinator creates this directory before the "
                                    "child starts."
                                ),
                            },
                            "depends_on": {
                                "type": "array",
                                "items": {"type": "string"},
                                "description": "Existing child-agent IDs that must finish before this task starts",
                            },
                        },
                        "required": ["name", "role", "task"],
                        "additionalProperties": False,
                    },
                },
            },
            {
                "type": "function",
                "function": {
                    "name": "send_agent_message",
                    "description": (
                        "Send a concise follow-up or correction to one running child agent."
                    ),
                    "parameters": {
                        "type": "object",
                        "properties": {
                            "agent_id": {"type": "string"},
                            "message": {"type": "string"},
                            "kind": {
                                "type": "string",
                                "enum": ["message", "question", "decision", "handoff", "blocker"],
                                "description": "Purpose of this concise collaboration message",
                            },
                            "references": {
                                "type": "array",
                                "items": {"type": "string"},
                                "description": "Relevant paths, artifact IDs, or interface names",
                            },
                        },
                        "required": ["agent_id", "message"],
                        "additionalProperties": False,
                    },
                },
            },
            {
                "type": "function",
                "function": {
                    "name": "publish_agent_artifact",
                    "description": (
                        "Publish a concise, reusable implementation handoff to the "
                        "team blackboard. This shares only the stated summary and "
                        "references, never private reasoning or full history."
                    ),
                    "parameters": {
                        "type": "object",
                        "properties": {
                            "title": {"type": "string"},
                            "summary": {"type": "string"},
                            "paths": {
                                "type": "array",
                                "items": {"type": "string"},
                            },
                            "recipient_agent_ids": {
                                "type": "array",
                                "items": {"type": "string"},
                                "description": "Only recipients that need an immediate notification",
                            },
                        },
                        "required": ["title", "summary"],
                        "additionalProperties": False,
                    },
                },
            },
            {
                "type": "function",
                "function": {
                    "name": "get_agent_collaboration",
                    "description": "Read the bounded public team blackboard: messages, handoffs, artifacts, dependencies, and file ownership.",
                    "parameters": {
                        "type": "object",
                        "properties": {},
                        "additionalProperties": False,
                    },
                },
            },
            {
                "type": "function",
                "function": {
                    "name": "wait_agents",
                    "description": (
                        "Wait for selected child agents and return the latest complete "
                        "team snapshot. Omit agent_ids to wait for every child."
                    ),
                    "parameters": {
                        "type": "object",
                        "properties": {
                            "agent_ids": {
                                "type": "array",
                                "items": {"type": "string"},
                            },
                            "timeout_ms": {
                                "type": "integer",
                                "minimum": 0,
                                "maximum": 600000,
                                "description": "Maximum wait in milliseconds",
                            },
                        },
                        "additionalProperties": False,
                    },
                },
            },
            {
                "type": "function",
                "function": {
                    "name": "list_agents",
                    "description": "Return the current complete child-agent team snapshot.",
                    "parameters": {
                        "type": "object",
                        "properties": {},
                        "additionalProperties": False,
                    },
                },
            },
            {
                "type": "function",
                "function": {
                    "name": "cancel_agent",
                    "description": "Request cooperative cancellation of one child agent.",
                    "parameters": {
                        "type": "object",
                        "properties": {"agent_id": {"type": "string"}},
                        "required": ["agent_id"],
                        "additionalProperties": False,
                    },
                },
            },
        ]

    def execute(
        self,
        tool_call: Dict[str, Any],
        conversation_id: Optional[str] = None,
        message_id: Optional[str] = None,
        runtime: Optional[Dict[str, Any]] = None,
    ) -> Any:
        """Execute a tool call"""
        tool_name = tool_call.get("tool")
        params = tool_call.get("params", {})

        scope_error = self._mutation_scope_error(tool_name, params, runtime)
        if scope_error:
            return scope_error

        stale_error = self._stale_mutation_error(tool_name, params)
        if stale_error:
            return stale_error

        if tool_name in MULTI_AGENT_TOOL_NAMES:
            return self._execute_multi_agent_tool(tool_name, params, runtime)

        if tool_name not in self.tools:
            return f"Error: Unknown tool '{tool_name}'"

        try:
            resolved_conversation_id = (
                conversation_id
                if conversation_id is not None
                else tool_call.get(
                    "conversation_id", getattr(self, "conversation_id", None)
                )
            )
            resolved_message_id = (
                message_id
                if message_id is not None
                else tool_call.get(
                    "message_id", getattr(self, "message_id", None)
                )
            )
            if tool_name == "update_plan":
                return self.execute_update_plan(
                    params,
                    conversation_id=resolved_conversation_id,
                    message_id=resolved_message_id,
                )
            if tool_name == "todo_write":
                return self.execute_todo_write(
                    params,
                    conversation_id=resolved_conversation_id,
                    message_id=resolved_message_id,
                )
            if tool_name == "view_image":
                return self.execute_view_image(
                    params,
                    conversation_id=resolved_conversation_id,
                    message_id=resolved_message_id,
                )
            if tool_name == "project_preview":
                return self.execute_project_preview(
                    params,
                    conversation_id=resolved_conversation_id,
                    message_id=resolved_message_id,
                )
            if tool_name == "load_skill":
                return self.execute_load_skill(
                    params, conversation_id=resolved_conversation_id,
                    message_id=resolved_message_id,
                )
            if tool_name in {"bash", "shell", "run_terminal_cmd", "pwsh"}:
                return self.execute_shell(params, runtime=runtime, tool_name=tool_name)
            if tool_name == "use_tool":
                return self.execute_use_tool(params, runtime=runtime)
            if tool_name in {"read", "file_read"}:
                return self.execute_file_read(
                    params,
                    conversation_id=resolved_conversation_id,
                    message_id=resolved_message_id,
                )
            if tool_name == "grep":
                return self.execute_grep(
                    params,
                    conversation_id=resolved_conversation_id,
                    message_id=resolved_message_id,
                )
            if tool_name == "glob":
                return self.execute_glob(
                    params,
                    conversation_id=resolved_conversation_id,
                    message_id=resolved_message_id,
                )
            result = self.tools[tool_name](params)
            return result
        except Exception as e:
            return f"Error executing {tool_name}: {str(e)}"

    def _mutation_scope_error(
        self,
        tool_name: object,
        params: object,
        runtime: Optional[Dict[str, Any]],
    ) -> str:
        """Reject mutations that escape a child agent's assigned write roots.

        Root task executors do not define a mutation scope and retain their
        existing behavior. Desktop child executors currently persist the scope
        on the executor instance, while other callers may provide the same
        policy in the per-call runtime.
        """
        raw_roots: object = _MUTATION_SCOPE_UNSET
        if isinstance(runtime, dict) and "mutation_scope_roots" in runtime:
            raw_roots = runtime["mutation_scope_roots"]
        elif hasattr(self, "mutation_scope_roots"):
            raw_roots = getattr(self, "mutation_scope_roots")
        if raw_roots is _MUTATION_SCOPE_UNSET:
            return ""

        normalized_name = str(tool_name or "").strip()
        if normalized_name in _SCOPED_COMMAND_TOOLS:
            return (
                "Error: command execution is unavailable while this agent has "
                "restricted write paths"
            )
        if normalized_name == "project_preview" and isinstance(params, dict):
            if str(params.get("action", "")).strip().lower() == "start":
                return (
                    "Error: starting a project preview is unavailable while this "
                    "agent has restricted write paths"
                )

        target_groups = _SCOPED_MUTATION_TARGETS.get(normalized_name)
        if target_groups is None:
            return ""

        roots = self._normalize_mutation_scope_roots(raw_roots)
        if not roots:
            return "Error: this agent has no allowed write paths"

        normalized_params = params if isinstance(params, dict) else {}
        targets = []
        for aliases in target_groups:
            target = next(
                (
                    normalized_params.get(key)
                    for key in aliases
                    if normalized_params.get(key) not in {None, ""}
                ),
                None,
            )
            if target is not None:
                targets.append(target)

        for target in targets:
            lexical_target, resolved_target = self._mutation_scope_paths(target)
            if not any(
                self._is_within_directory(lexical_target, lexical_root)
                and self._is_within_directory(resolved_target, resolved_root)
                for lexical_root, resolved_root in roots
            ):
                return (
                    "Error: mutation target is outside this agent's allowed "
                    f"write paths: {target}"
                )
        return ""

    def _stale_mutation_error(
        self, tool_name: object, params: object
    ) -> str:
        """Reject edits whose target changed since the last read.

        A successful read records the target's version fingerprint; if the
        file changed on disk before a later patch, the model is editing
        content it never actually saw, so the mutation fails and demands a
        re-read instead of silently corrupting the new content. Whole-file
        writes replace everything and are exempt from this check.
        """
        if not getattr(self, "_read_versions", None):
            return ""
        normalized_name = str(tool_name or "").strip()
        target_groups = _STALE_CHECKED_TARGETS.get(normalized_name)
        if target_groups is None:
            return ""
        normalized_params = params if isinstance(params, dict) else {}
        for aliases in target_groups:
            target = next(
                (
                    normalized_params.get(key)
                    for key in aliases
                    if isinstance(normalized_params.get(key), str)
                    and normalized_params.get(key)
                ),
                None,
            )
            if target is None:
                continue
            path = os.path.normcase(
                os.path.realpath(FileTool.expand_path(target))
            )
            recorded = self._read_versions.get(path)
            if recorded is None:
                return ""
            current = self._stat_version(path)
            if current is not None and current != recorded:
                return (
                    f"Error: {target} changed since it was read; the edit target "
                    "is stale. Re-read the file before editing it again."
                )
            return ""
        return ""

    @staticmethod
    def _stat_version(path: str) -> Optional[Tuple[int, int]]:
        """Fingerprint a file's current on-disk version (mtime_ns, size)."""
        try:
            stat = os.stat(path)
            return (stat.st_mtime_ns, stat.st_size)
        except OSError:
            return None

    def _note_read_version(self, file_path: str) -> None:
        """Record a successful read's version fingerprint for stale checks."""
        try:
            path = os.path.normcase(
                os.path.realpath(FileTool.expand_path(file_path))
            )
            if os.path.isfile(path):
                version = self._stat_version(path)
                if version is not None:
                    self._read_versions[path] = version
        except Exception:
            pass

    def _normalize_mutation_scope_roots(
        self, raw_roots: object
    ) -> tuple[tuple[Path, Path], ...]:
        """Return distinct lexical/resolved path pairs for one write scope."""
        if isinstance(raw_roots, (str, os.PathLike)):
            candidates = (raw_roots,)
        else:
            try:
                candidates = tuple(raw_roots)  # type: ignore[arg-type]
            except TypeError:
                candidates = ()

        roots = []
        seen = set()
        for candidate in candidates:
            if candidate in {None, ""}:
                continue
            candidate_path = Path(candidate)
            while candidate_path.name in {"*", "**"}:
                candidate_path = candidate_path.parent
            if any(marker in str(candidate_path) for marker in ("*", "?", "[")):
                continue
            lexical_path, resolved_path = self._mutation_scope_paths(candidate_path)
            pair = (lexical_path, resolved_path)
            if pair not in seen:
                seen.add(pair)
                roots.append(pair)
        return tuple(roots)

    def _mutation_scope_paths(self, path: object) -> tuple[Path, Path]:
        """Resolve one path lexically and through any existing symlinks."""
        value = str(path).strip()
        value = value.replace("桌面", "Desktop")
        value = value.replace("文档", "Documents")
        value = value.replace("下载", "Downloads")
        candidate = Path(os.path.expanduser(value))
        if not candidate.is_absolute():
            candidate = self.project_root / candidate
        lexical_path = Path(os.path.abspath(str(candidate)))
        return lexical_path, candidate.resolve(strict=False)

    @staticmethod
    def _execute_multi_agent_tool(
        tool_name: str,
        params: Dict[str, Any],
        runtime: Optional[Dict[str, Any]],
    ) -> Any:
        """Dispatch collaboration tools only through the active graph runtime."""
        dispatch = (runtime or {}).get("multi_agent_dispatch")
        if not callable(dispatch):
            return (
                f"Error: {tool_name} is unavailable because Multi-Agent Mode "
                "is not active"
            )
        try:
            result = dispatch(tool_name, dict(params or {}))
        except Exception as exc:
            return f"Error executing {tool_name}: {exc}"
        if isinstance(result, (str, ToolExecutionResult)):
            return result
        return json.dumps(result, ensure_ascii=False, default=str)

    @staticmethod
    def _plan_key(
        conversation_id: Optional[str], message_id: Optional[str]
    ) -> str:
        conversation = str(conversation_id or "").strip()
        message = str(message_id or "").strip()
        return f"{conversation}:{message}" if conversation or message else "default"

    def _plan_tool_for(
        self, conversation_id: Optional[str], message_id: Optional[str]
    ) -> PlanTool:
        key = self._plan_key(conversation_id, message_id)
        with self._plan_tools_lock:
            return self._plan_tools.setdefault(key, PlanTool())

    def execute_update_plan(
        self,
        params: Dict[str, Any],
        conversation_id: Optional[str] = None,
        message_id: Optional[str] = None,
    ) -> str:
        """Replace the plan isolated to the active task message."""
        return self._plan_tool_for(conversation_id, message_id).update(params)

    def get_plan_snapshot(
        self,
        conversation_id: Optional[str] = None,
        message_id: Optional[str] = None,
    ) -> Optional[Dict[str, Any]]:
        """Return the latest snapshot for one task without creating it."""
        key = self._plan_key(conversation_id, message_id)
        with self._plan_tools_lock:
            tool = self._plan_tools.get(key)
        return tool.snapshot() if tool else None

    def discard_plan_snapshot(
        self,
        conversation_id: Optional[str] = None,
        message_id: Optional[str] = None,
    ) -> None:
        """Release one completed task's legacy and todo plan state."""
        key = self._plan_key(conversation_id, message_id)
        with self._plan_tools_lock:
            self._plan_tools.pop(key, None)
        with self._todo_state_lock:
            self._todo_states.pop(key, None)
        with self._loaded_skills_lock:
            self._loaded_skills.pop(key, None)

    def clear_plan_snapshots(self, conversation_id: Optional[str] = None) -> None:
        """Release legacy and todo plan state for one conversation."""
        conversation = str(conversation_id or "").strip()
        with self._plan_tools_lock:
            if not conversation:
                self._plan_tools.clear()
            else:
                prefix = f"{conversation}:"
                for key in tuple(self._plan_tools):
                    if key == conversation or key.startswith(prefix):
                        self._plan_tools.pop(key, None)
        with self._todo_state_lock:
            if not conversation:
                self._todo_states.clear()
            else:
                prefix = f"{conversation}:"
                for key in tuple(self._todo_states):
                    if key == conversation or key.startswith(prefix):
                        self._todo_states.pop(key, None)
        with self._loaded_skills_lock:
            if not conversation:
                self._loaded_skills.clear()
            else:
                prefix = f"{conversation}:"
                for key in tuple(self._loaded_skills):
                    if key == conversation or key.startswith(prefix):
                        self._loaded_skills.pop(key, None)

    def get_todo_snapshot(
        self,
        conversation_id: Optional[str] = None,
        message_id: Optional[str] = None,
    ) -> list[Dict[str, Any]]:
        """Return the current Grok-style todos without mutating their state."""
        key = self._plan_key(conversation_id, message_id)
        with self._todo_state_lock:
            state = self._todo_states.get(key, {})
            return [dict(item) for item in state.get("todos", {}).values()]

    def get_running_background_tasks(self) -> list[Dict[str, str]]:
        """Return active background commands for compaction reminders."""
        with self._background_tasks_lock:
            return [
                {"task_id": task_id, "command": str(task.get("command", ""))}
                for task_id, task in self._background_tasks.items()
                if task.get("process") is not None and task["process"].poll() is None
            ]

    def get_loaded_skills(
        self,
        conversation_id: Optional[str] = None,
        message_id: Optional[str] = None,
    ) -> list[str]:
        """Return skills successfully loaded for the active task."""
        key = self._plan_key(conversation_id, message_id)
        with self._loaded_skills_lock:
            return sorted(self._loaded_skills.get(key, set()))

    @staticmethod
    def _task_image_key(
        conversation_id: Optional[str], message_id: Optional[str]
    ) -> str:
        conversation = str(conversation_id or "").strip()
        message = str(message_id or "").strip()
        if not conversation or not message:
            raise ValueError("Current task image context is unavailable")
        return f"{conversation}:{message}"

    def register_task_images(
        self,
        conversation_id: Optional[str],
        message_id: Optional[str],
        images: List[Dict[str, Any]],
    ) -> None:
        """Register validated image files for exactly one active task run."""
        key = self._task_image_key(conversation_id, message_id)
        records: Dict[str, _TaskImage] = {}
        for image in images or []:
            if not isinstance(image, dict):
                continue
            raw_path = str(image.get("path", "")).strip()
            mime_type = str(image.get("type", "")).strip().lower()
            if not raw_path or mime_type not in _TASK_IMAGE_MIME_TYPES:
                continue
            try:
                path = Path(raw_path).expanduser().resolve(strict=True)
                content = path.read_bytes()
            except OSError as exc:
                raise ValueError(f"Task image is unavailable: {raw_path}") from exc
            size = len(content)
            try:
                expected_size = int(image.get("size", size) or 0)
            except (TypeError, ValueError) as exc:
                raise ValueError("Task image has an invalid size") from exc
            if not size or size > MAX_TASK_IMAGE_BYTES or expected_size != size:
                raise ValueError("Task image has an invalid size")
            records[str(path)] = _TaskImage(
                path=path,
                name=Path(str(image.get("name", path.name))).name or path.name,
                mime_type=mime_type,
                size=size,
                digest=hashlib.sha256(content).hexdigest(),
            )
        with self._task_images_lock:
            if records:
                self._task_images[key] = records
            else:
                self._task_images.pop(key, None)

    def clear_task_images(
        self, conversation_id: Optional[str], message_id: Optional[str]
    ) -> None:
        """Forget the in-memory allowlist once its task run has finished."""
        try:
            key = self._task_image_key(conversation_id, message_id)
        except ValueError:
            return
        with self._task_images_lock:
            self._task_images.pop(key, None)

    def register_task_reference_roots(
        self,
        conversation_id: Optional[str],
        message_id: Optional[str],
        paths: List[str],
    ) -> None:
        """Remember dropped folders as explicit task context."""
        key = self._task_image_key(conversation_id, message_id)
        roots = []
        seen = set()
        for raw_path in paths or []:
            path = Path(str(raw_path or "")).expanduser().resolve(strict=True)
            if not path.is_dir():
                raise ValueError(f"Reference folder is unavailable: {raw_path}")
            if path in seen:
                continue
            seen.add(path)
            roots.append(path)
        with self._task_reference_roots_lock:
            if roots:
                self._task_reference_roots[key] = tuple(roots)
            else:
                self._task_reference_roots.pop(key, None)

    def clear_task_reference_roots(
        self, conversation_id: Optional[str], message_id: Optional[str]
    ) -> None:
        """Forget dropped reference folders after their task run finishes."""
        try:
            key = self._task_image_key(conversation_id, message_id)
        except ValueError:
            return
        with self._task_reference_roots_lock:
            self._task_reference_roots.pop(key, None)

    def _task_reference_root_for_path(
        self,
        path: Path,
        conversation_id: Optional[str],
        message_id: Optional[str],
    ) -> Optional[Path]:
        try:
            key = self._task_image_key(conversation_id, message_id)
        except ValueError:
            return None
        with self._task_reference_roots_lock:
            roots = self._task_reference_roots.get(key, ())
        return next(
            (root for root in roots if self._is_within_directory(path, root)),
            None,
        )

    @staticmethod
    def _is_within_directory(path: Path, directory: Path) -> bool:
        """Return whether a resolved path remains inside one allowlisted folder."""
        try:
            path.relative_to(directory)
        except ValueError:
            return False
        return True

    def _resolve_project_path(self, path: object, default: str = ".") -> str:
        """Resolve relative tool paths against this task's bound project root."""
        raw_path = str(path if path not in {None, ""} else default).strip() or default
        raw_path = raw_path.replace("桌面", "Desktop")
        raw_path = raw_path.replace("文档", "Documents")
        raw_path = raw_path.replace("下载", "Downloads")
        candidate = Path(os.path.expanduser(raw_path))
        if not candidate.is_absolute():
            candidate = self.project_root / candidate
        return str(candidate.resolve(strict=False))

    def _mutation_path_error(self, *paths: object) -> str:
        """Reject direct file-tool mutations of protected OS-Agent source."""
        for raw_path in paths:
            if raw_path in {None, ""}:
                continue
            value = str(raw_path).strip()
            value = value.replace("桌面", "Desktop")
            value = value.replace("文档", "Documents")
            value = value.replace("下载", "Downloads")
            candidate = Path(os.path.expanduser(value))
            if not candidate.is_absolute():
                candidate = self.project_root / candidate
            lexical_path = Path(os.path.abspath(str(candidate)))
            resolved_path = candidate.resolve(strict=False)

            lexical_protected = any(
                self._is_within_directory(lexical_path, region)
                for region in self._protected_regions
            )
            resolved_protected = any(
                self._is_within_directory(resolved_path, region)
                for region in self._protected_regions
            )
            lexical_allowed = any(
                self._is_within_directory(lexical_path, root)
                for root in self.protected_write_roots
            )
            resolved_allowed = any(
                self._is_within_directory(resolved_path, root)
                for root in self.protected_write_roots
            )
            if (lexical_protected and not lexical_allowed) or (
                resolved_protected and not resolved_allowed
            ):
                return (
                    "Error: JCodex source is protected; only "
                    "workspace/output and workspace/temp may be modified"
                )
        return ""

    def _workspace_image_source(self, path: Path) -> Optional[str]:
        """Identify the allowlisted workspace image folder containing ``path``."""
        if self._is_within_directory(path, self.workspace_temp_root):
            return "workspace temp"
        if self._is_within_directory(path, self.workspace_output_root):
            return "workspace output"
        return None

    @staticmethod
    def _image_tool_result(
        path: Path, mime_type: str, content: bytes, source: str
    ) -> ToolExecutionResult:
        """Build the transient model-only image payload without persisting bytes."""
        encoded = base64.b64encode(content).decode("ascii")
        return ToolExecutionResult(
            content=(
                f"Loaded {source}: {path}. "
                "The image is available to the model for this task turn."
            ),
            model_inputs=[
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:{mime_type};base64,{encoded}"
                    },
                }
            ],
        )

    def _load_workspace_image(
        self, path: Path, source: str
    ) -> ToolExecutionResult | str:
        """Load a supported image from one explicitly allowlisted workspace folder."""
        try:
            if not path.is_file():
                return f"Error: {source} image path is not a file"
            if path.stat().st_size > MAX_TASK_IMAGE_BYTES:
                return f"Error: {source} image exceeds the 12 MB limit"
            content = path.read_bytes()
        except OSError:
            return f"Error: {source} image is unavailable"
        if not content:
            return f"Error: {source} image is empty"
        mime_type = _detect_supported_image_mime(content)
        if mime_type is None:
            return f"Error: {source} image must be PNG, JPEG, or WebP"
        return self._image_tool_result(path, mime_type, content, f"{source} image")

    def execute_view_image(
        self,
        params: Dict[str, Any],
        conversation_id: Optional[str] = None,
        message_id: Optional[str] = None,
    ) -> ToolExecutionResult | str:
        """Load one allowlisted image and return it only to the active model run."""
        requested_path = str(params.get("path", "")).strip()
        if not requested_path:
            return "Error: path parameter required"
        try:
            key = self._task_image_key(conversation_id, message_id)
            path = Path(requested_path).expanduser().resolve(strict=True)
        except (OSError, ValueError):
            return "Error: image path is invalid or unavailable"
        workspace_source = self._workspace_image_source(path)
        if workspace_source:
            return self._load_workspace_image(path, workspace_source)
        with self._task_images_lock:
            record = self._task_images.get(key, {}).get(str(path))
        if record is None:
            return (
                "Error: only current-conversation image attachments or images "
                "inside workspace/temp or workspace/output may be viewed"
            )
        try:
            content = record.path.read_bytes()
        except OSError:
            return "Error: image attachment is no longer available"
        if (
            len(content) != record.size
            or hashlib.sha256(content).hexdigest() != record.digest
        ):
            return "Error: image attachment changed after it was registered"
        return self._image_tool_result(
            record.path,
            record.mime_type,
            content,
            "current-conversation image attachment",
        )

    def execute_shell(
        self,
        params: Dict[str, Any],
        runtime: Optional[Dict[str, Any]] = None,
        tool_name: Optional[str] = None,
    ) -> str:
        """Execute shell command.

        Args:
            params: Tool parameters (command, workdir, timeout, ...).
            runtime: Per-call runtime (cancellation signals).
            tool_name: Dispatching tool name. ``pwsh`` runs the command through
                PowerShell; other names use the platform shell.
        """
        command = params.get("command", "")
        if not command:
            return "Error: command parameter required"

        # The pwsh tool advertises the PowerShell dialect, so its commands must
        # run inside PowerShell, not cmd.exe (which shell=True would use).
        dialect = "pwsh" if tool_name == "pwsh" else "system"
        if bool(params.get("is_background", False)):
            return self._start_background_command(params, dialect=dialect)

        timeout = params.get("timeout")
        if timeout is not None:
            try:
                timeout = max(float(timeout) / 1000, 0.1)
            except (TypeError, ValueError):
                return "Error: timeout must be a number of milliseconds"

        result = self.shell_tool.execute(
            command,
            cwd=(
                self._resolve_project_path(params.get("workdir"), ".")
                if hasattr(self, "project_root")
                else params.get("workdir")
            ),
            timeout=timeout,
            cancel_event=(runtime or {}).get("cancel_event"),
            cancelled=(runtime or {}).get("cancelled"),
            dialect=dialect,
        )
        return self.shell_tool.format_result(result)

    def execute_file_read(
        self,
        params: Dict[str, Any],
        conversation_id: Optional[str] = None,
        message_id: Optional[str] = None,
    ) -> str:
        """Read file with pagination support - aligned with OpenCode"""
        file_path = params.get("filePath", "") or params.get("path", "")
        if not file_path:
            return "Error: filePath parameter required"

        file_path = self._resolve_project_path(file_path)
        resolved_path = Path(file_path)
        if not resolved_path.exists():
            return "Error: reference path is unavailable"

        suffix = os.path.splitext(str(file_path))[1].lower()
        if suffix in {".pdf", ".doc", ".docx"}:
            return self.execute_read_pdf({"path": file_path})
        if suffix in {".xlsx", ".xlsm"}:
            return self.execute_read_excel({"path": file_path})

        offset = params.get("offset")
        limit = params.get("limit")
        if offset is None and params.get("from") is not None:
            offset = int(params.get("from", 0)) + 1
        if limit is None and params.get("lines") is not None:
            limit = params.get("lines")

        success, content = self.file_tool.read_file(file_path, offset, limit)

        if success:
            self._note_read_version(file_path)
            return content
        return f"Error: {content}"

    def get_glob_tool_definition(self) -> Dict[str, Any]:
        """Expose LangChain's scoped file search through the existing glob API."""
        return {
            "type": "function",
            "function": {
                "name": "glob",
                "description": (
                    "Find files by shell-style filename pattern beneath the project "
                    "root. Use path to scope the search; matches are returned as "
                    "relative paths."
                ),
                "parameters": {
                    "type": "object",
                    "properties": {
                        "pattern": {
                            "type": "string",
                            "description": "Filename pattern such as '*.py' or 'package.json'. Prefix with '**/' to include subdirectories.",
                        },
                        "path": {
                            "type": "string",
                            "description": "Project-relative directory to search (defaults to project root)",
                        },
                    },
                    "required": ["pattern"],
                },
            },
        }

    def execute_glob(
        self,
        params: Dict[str, Any],
        conversation_id: Optional[str] = None,
        message_id: Optional[str] = None,
    ) -> str:
        """Delegate filename matching to LangChain's scoped FileSearchTool."""
        pattern = str(params.get("pattern", "")).strip()
        path = str(params.get("path", ".")).strip() or "."
        if not pattern:
            return "Error: pattern parameter required"

        parts = [part for part in pattern.replace("\\", "/").split("/") if part]
        if not parts:
            return "Error: pattern parameter required"
        filename_pattern = parts[-1]
        directory_parts = [part for part in parts[:-1] if part not in {".", "**"}]
        if any(any(char in part for char in "*?[") for part in directory_parts):
            return (
                "Error: directory wildcards are unsupported; pass the directory "
                "with path and use a filename pattern"
            )
        recursive = "**" in parts
        search_path = str(Path(path).joinpath(*directory_parts))
        resolved_search_path = Path(self._resolve_project_path(search_path))
        reference_root = self._task_reference_root_for_path(
            resolved_search_path, conversation_id, message_id
        )
        if reference_root is not None:
            try:
                matches = [
                    candidate.relative_to(resolved_search_path).as_posix()
                    for candidate in resolved_search_path.rglob(filename_pattern)
                    if candidate.is_file()
                ]
            except OSError as exc:
                return f"Error: file search failed: {exc}"
            if not recursive:
                matches = [match for match in matches if "/" not in match]
            if directory_parts:
                prefix = Path(*directory_parts)
                matches = [str(prefix / match) for match in matches]
            return "\n".join(sorted(matches)) or (
                f"No files found for pattern {pattern} in directory {path}"
            )
        if not self._is_within_directory(resolved_search_path, self.project_root):
            try:
                matches = [
                    candidate.relative_to(resolved_search_path).as_posix()
                    for candidate in resolved_search_path.rglob(filename_pattern)
                    if candidate.is_file()
                ]
            except OSError as exc:
                return f"Error: file search failed: {exc}"
            if not recursive:
                matches = [match for match in matches if "/" not in match]
            if directory_parts:
                prefix = Path(*directory_parts)
                matches = [str(prefix / match) for match in matches]
            return "\n".join(sorted(matches)) or (
                f"No files found for pattern {pattern} in directory {path}"
            )
        try:
            result = self._file_search_tool.invoke(
                {"pattern": filename_pattern, "dir_path": search_path}
            )
        except Exception as exc:
            return f"Error: file search failed: {exc}"
        output = str(result)
        if output.startswith("Error:"):
            return output
        matches = [line for line in output.splitlines() if line.strip()]
        if output.startswith("No files found"):
            return f"No files found for pattern {pattern} in directory {path}"
        if not recursive:
            matches = [match for match in matches if "/" not in match]
        if directory_parts:
            prefix = Path(*directory_parts)
            matches = [str(prefix / match) for match in matches]
        if not matches:
            return f"No files found for pattern {pattern} in directory {path}"
        return "\n".join(matches)

    def execute_grep(
        self,
        params: Dict[str, Any],
        conversation_id: Optional[str] = None,
        message_id: Optional[str] = None,
    ) -> str:
        """Search content with relative paths rooted at the bound project."""
        normalized = dict(params)
        normalized["path"] = self._resolve_project_path(params.get("path"), ".")
        return execute_grep(normalized)

    def execute_edit(self, params: Dict[str, Any]) -> str:
        """Edit a file with relative paths rooted at the bound project."""
        normalized = dict(params)
        path = params.get("filePath") or params.get("file_path") or params.get("path")
        protection_error = self._mutation_path_error(path)
        if protection_error:
            return protection_error
        normalized["filePath"] = self._resolve_project_path(path)
        normalized["oldString"] = params.get("oldString", params.get("old_string", ""))
        normalized["newString"] = params.get("newString", params.get("new_string", ""))
        normalized["replaceAll"] = params.get("replaceAll", params.get("replace_all", False))
        result = execute_edit(normalized)
        if not str(result).startswith("Error:"):
            # The mutation succeeded: refresh the recorded version so a later
            # edit of this file is not mistaken for a stale target.
            self._note_read_version(normalized["filePath"])
        return result

    def execute_read_excel(self, params: Dict[str, Any]) -> str:
        """Read an Excel workbook as tab-separated sheet content."""
        path = params.get("path", "") or params.get("filePath", "")
        if not path:
            return "Error: path parameter required"
        try:
            import openpyxl

            workbook = openpyxl.load_workbook(
                self._resolve_project_path(path), read_only=True, data_only=True
            )
            lines = []
            for sheet in workbook.worksheets:
                lines.append(f"## Sheet: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    values = ["" if value is None else str(value) for value in row]
                    if any(values):
                        lines.append("\t".join(values))
                    if sum(len(line) + 1 for line in lines) >= 50000:
                        lines.append("[Workbook content truncated]")
                        break
                if lines and lines[-1] == "[Workbook content truncated]":
                    break
            workbook.close()
            return "Excel contents:\n" + "\n".join(lines)
        except ImportError:
            return "Error: openpyxl not installed. Try: pip install openpyxl"
        except Exception as exc:
            return f"Error reading Excel workbook: {str(exc)}"

    def execute_file_write(self, params: Dict[str, Any]) -> str:
        """Write file"""
        path = params.get("path", "")
        content = params.get("content", "")
        if not path:
            return "Error: path parameter required"

        protection_error = self._mutation_path_error(path)
        if protection_error:
            return protection_error
        path = self._resolve_project_path(path)
        success, message = self.file_tool.write_file(path, content)
        if success:
            # Refresh the recorded version so a later edit of this file is not
            # mistaken for a stale target.
            self._note_read_version(path)
            return message
        return f"Error: {message}"

    def execute_file_list(self, params: Dict[str, Any]) -> str:
        """List files"""
        path = self._resolve_project_path(
            params.get("path") or params.get("target_directory"), "."
        )
        success, files = self.file_tool.list_files(path)

        if success:
            if not files:
                return f"Files in {path}:\nnone"
            file_list = "\n".join(files)
            return f"Files in {path}:\n{file_list}"
        return f"Error: {files[0] if files else 'Unknown error'}"

    def execute_todo_write(
        self,
        params: Dict[str, Any],
        conversation_id: Optional[str] = None,
        message_id: Optional[str] = None,
    ) -> str:
        """Persist one task's Grok-style todo plan and return its snapshot."""
        todos = params.get("todos")
        if not isinstance(todos, list):
            return "Error: todos parameter must be an array"
        key = self._plan_key(conversation_id, message_id)
        with self._todo_state_lock:
            state = self._todo_states.setdefault(
                key, {"version": 0, "todos": {}}
            )
            current = dict(state["todos"])
            if not bool(params.get("merge", True)):
                current.clear()
            for index, item in enumerate(todos):
                if not isinstance(item, dict) or not str(item.get("id", "")).strip():
                    return f"Error: todos[{index}].id is required"
                identifier = str(item["id"]).strip()
                previous = current.get(identifier, {})
                content = str(
                    item.get("content", previous.get("content", identifier))
                ).strip()
                status = str(
                    item.get("status", previous.get("status", "pending"))
                ).strip()
                if status not in {"pending", "in_progress", "completed", "cancelled"}:
                    return f"Error: invalid todo status: {status}"
                current[identifier] = {
                    "id": identifier,
                    "content": content or identifier,
                    "status": status,
                }
            if sum(
                item["status"] == "in_progress" for item in current.values()
            ) > 1:
                return "Error: only one todo may be in_progress"
            if not current:
                return "Error: todo plan must contain at least one item"
            state["version"] = int(state["version"]) + 1
            state["todos"] = current
            snapshot = list(current.values())
            version = int(state["version"])
        active = next(
            (item["content"] for item in snapshot if item["status"] == "in_progress"),
            "",
        )
        return json.dumps(
            {
                "success": True,
                "version": version,
                "completed": sum(
                    item["status"] in {"completed", "cancelled"}
                    for item in snapshot
                ),
                "total": len(snapshot),
                "current_step": active,
                "todos": snapshot,
            },
            ensure_ascii=False,
        )

    def execute_memory_search(self, params: Dict[str, Any]) -> str:
        """Search the bound Grok-style memory store."""
        if self.memory_store is None:
            return "Memory is not enabled."
        query = str(params.get("query", "")).strip()
        if not query:
            return "Error: query parameter required"
        results = self.memory_store.search(
            query,
            limit=max(1, int(params.get("max_results", 6))),
            min_score=float(params.get("min_score", self.memory_store.min_score)),
        )
        if not results:
            return "No memory results found for query."
        return self.memory_store.format_memory_context(results)

    def execute_memory_get(self, params: Dict[str, Any]) -> str:
        """Read one memory Markdown file with Grok-compatible line bounds."""
        if self.memory_store is None:
            return "Memory is not enabled."
        requested = str(params.get("path", "")).strip()
        if not requested:
            return "Error: path parameter required"
        path = (self.memory_store.global_dir / requested).resolve()
        try:
            path.relative_to(self.memory_store.global_dir)
        except ValueError:
            return "Error: memory path is outside the memory root"
        if not path.is_file() or path.suffix.lower() != ".md":
            return "Error: memory file not found"
        start = max(0, int(params.get("from", 0)))
        raw_lines = path.read_text(encoding="utf-8", errors="replace").split("\n")
        count = params.get("lines")
        selected = (
            raw_lines[start:]
            if count is None
            else raw_lines[start : start + max(0, int(count))]
        )
        numbered = "\n".join(
            f"{start + index + 1}→{line}" for index, line in enumerate(selected)
        )
        return f"**File:** {requested}\n**Lines:** {len(selected)}\n\n{numbered}"

    def _start_background_command(
        self, params: Dict[str, Any], dialect: str = "system"
    ) -> str:
        command = str(params.get("command", "")).strip()
        if not command:
            return "Error: command parameter required"
        workdir = self._resolve_project_path(params.get("workdir"), ".")
        self.task_log_root.mkdir(parents=True, exist_ok=True)
        task_id = uuid.uuid4().hex[:12]
        output_path = self.task_log_root / f"task-{task_id}.log"
        output_handle = output_path.open("wb")
        options: Dict[str, Any] = {}
        if os.name == "posix":
            options["start_new_session"] = True
        elif os.name == "nt":
            options["creationflags"] = subprocess.CREATE_NEW_PROCESS_GROUP
        args: object = command
        if dialect == "pwsh":
            pwsh = ShellTool._resolve_pwsh()
            if pwsh is None:
                output_handle.close()
                return (
                    "Error: PowerShell (pwsh or powershell) was not found on PATH; "
                    "the pwsh tool requires it."
                )
            args = [pwsh, "-NoProfile", "-NonInteractive", "-Command", command]
        try:
            process = subprocess.Popen(
                args,
                shell=(dialect != "pwsh"),
                cwd=workdir,
                stdout=output_handle,
                stderr=subprocess.STDOUT,
                **options,
            )
        except Exception as e:
            output_handle.close()
            return f"Error: failed to start background command: {e}"
        with self._background_tasks_lock:
            self._background_tasks[task_id] = {
                "process": process,
                "command": command,
                "output_path": output_path,
                "output_handle": output_handle,
                "started_at": time.time(),
            }
        return json.dumps(
            {
                "task_id": task_id,
                "status": "running",
                "output_file": str(output_path),
            }
        )

    def execute_monitor(self, params: Dict[str, Any]) -> str:
        normalized = dict(params)
        normalized["is_background"] = True
        return self._start_background_command(normalized)

    def execute_get_task_output(self, params: Dict[str, Any]) -> str:
        task_ids = params.get("task_ids")
        if isinstance(task_ids, str):
            task_ids = [task_ids]
        if not isinstance(task_ids, list) or not task_ids:
            return "Error: task_ids parameter must be a non-empty array"
        timeout_ms = max(0, min(600000, int(params.get("timeout_ms", 0))))
        deadline = time.monotonic() + timeout_ms / 1000
        while timeout_ms:
            with self._background_tasks_lock:
                running = any(
                    task_id in self._background_tasks
                    and self._background_tasks[task_id]["process"].poll() is None
                    for task_id in task_ids
                )
            if not running or time.monotonic() >= deadline:
                break
            time.sleep(0.05)
        results = []
        with self._background_tasks_lock:
            for task_id in task_ids:
                task = self._background_tasks.get(str(task_id))
                if task is None:
                    results.append({"task_id": task_id, "status": "not_found"})
                    continue
                process = task["process"]
                code = process.poll()
                if code is not None and not task["output_handle"].closed:
                    task["output_handle"].close()
                try:
                    # The log is written as raw bytes; decode with the same
                    # UTF-8 -> GBK -> Latin-1 chain as foreground output so
                    # Windows cmd/PowerShell (cp936) output stays readable.
                    raw = task["output_path"].read_bytes()[-48000:]
                    output = decode_stream(raw)[-12000:]
                except OSError:
                    output = ""
                results.append(
                    {
                        "task_id": task_id,
                        "status": "running" if code is None else "completed",
                        "exit_code": code,
                        "output": output,
                        "output_file": str(task["output_path"]),
                    }
                )
        return json.dumps({"results": results}, ensure_ascii=False)

    def execute_kill_task(self, params: Dict[str, Any]) -> str:
        task_id = str(params.get("task_id", "")).strip()
        if not task_id:
            return "Error: task_id parameter required"
        with self._background_tasks_lock:
            task = self._background_tasks.get(task_id)
        if task is None:
            return json.dumps({"success": False, "task_id": task_id})
        process = task["process"]
        if process.poll() is None:
            try:
                if os.name == "posix":
                    os.killpg(process.pid, signal.SIGTERM)
                else:
                    process.terminate()
            except OSError:
                pass
        return json.dumps({"success": True, "task_id": task_id})

    def execute_search_tool(self, params: Dict[str, Any]) -> str:
        query = str(params.get("query", "")).lower().strip()
        if not query:
            return "Error: query parameter required"
        matches = []
        for definition in self.get_available_tools():
            function = definition.get("function", {})
            haystack = (
                f"{function.get('name', '')} {function.get('description', '')}".lower()
            )
            if all(term in haystack for term in query.split()):
                matches.append(function)
        return json.dumps(
            {"status": "complete", "tools": matches[:12]}, ensure_ascii=False
        )

    def execute_use_tool(
        self,
        params: Dict[str, Any],
        runtime: Optional[Dict[str, Any]] = None,
    ) -> Any:
        """Invoke one tool while preserving the caller's runtime policy."""
        tool_name = str(params.get("tool_name", "")).strip()
        tool_input = params.get("tool_input", {})
        if tool_name in {"use_tool", "search_tool"}:
            return "Error: use_tool cannot recursively invoke tool routers"
        if tool_name not in self.tools:
            return f"Error: Tool not found: {tool_name}"
        if not isinstance(tool_input, dict):
            return "Error: tool_input must be an object"
        return self.execute(
            {"tool": tool_name, "params": tool_input}, runtime=runtime
        )

    @staticmethod
    def _parse_scheduler_interval(value: str) -> int:
        match = re.fullmatch(r"\s*(\d+)\s*([smhd])\s*", str(value).lower())
        if not match:
            raise ValueError("interval must use syntax such as 60s, 5m, 2h, or 1d")
        amount = int(match.group(1))
        seconds = amount * {"s": 1, "m": 60, "h": 3600, "d": 86400}[match.group(2)]
        if seconds < 60:
            raise ValueError("scheduled intervals must be at least 60 seconds")
        return seconds

    def _arm_scheduled_task(self, task_id: str, delay: float) -> None:
        timer = threading.Timer(max(0.0, delay), self._fire_scheduled_task, [task_id])
        timer.daemon = True
        with self._scheduled_tasks_lock:
            task = self._scheduled_tasks.get(task_id)
            if task is None:
                return
            task["timer"] = timer
            task["next_fire"] = time.time() + max(0.0, delay)
        timer.start()

    def _fire_scheduled_task(self, task_id: str) -> None:
        with self._scheduled_tasks_lock:
            task = self._scheduled_tasks.get(task_id)
            if task is None:
                return
            prompt = task["prompt"]
            recurring = task["recurring"]
            interval_seconds = task["interval_seconds"]
            task["last_fire"] = time.time()
        callback = self.scheduled_prompt_callback
        if callback is not None:
            try:
                callback(task_id, prompt)
            except Exception as exc:
                with self._scheduled_tasks_lock:
                    if task_id in self._scheduled_tasks:
                        self._scheduled_tasks[task_id]["last_error"] = str(exc)
        if recurring:
            self._arm_scheduled_task(task_id, interval_seconds)
        else:
            with self._scheduled_tasks_lock:
                self._scheduled_tasks.pop(task_id, None)

    def execute_scheduler_create(self, params: Dict[str, Any]) -> str:
        task_id = str(params.get("task_id", "")).strip()
        updating = bool(task_id)
        with self._scheduled_tasks_lock:
            existing = self._scheduled_tasks.get(task_id) if updating else None
        if updating and existing is None:
            return f"Error: scheduled task not found: {task_id}"
        interval = params.get("interval", existing.get("interval") if existing else None)
        prompt = params.get("prompt", existing.get("prompt") if existing else None)
        if not interval or not str(prompt or "").strip():
            return "Error: interval and prompt are required when creating a task"
        try:
            interval_seconds = self._parse_scheduler_interval(str(interval))
        except ValueError as exc:
            return f"Error: {exc}"
        if not task_id:
            task_id = uuid.uuid4().hex[:12]
        if existing and existing.get("timer"):
            existing["timer"].cancel()
        task = {
            "id": task_id,
            "interval": str(interval),
            "interval_seconds": interval_seconds,
            "prompt": str(prompt).strip(),
            "recurring": bool(params.get("recurring", existing.get("recurring", True) if existing else True)),
            "durable": bool(params.get("durable", existing.get("durable", False) if existing else False)),
            "foreground": bool(params.get("foreground", existing.get("foreground", False) if existing else False)),
            "created_at": existing.get("created_at", time.time()) if existing else time.time(),
            "last_fire": existing.get("last_fire") if existing else None,
        }
        with self._scheduled_tasks_lock:
            self._scheduled_tasks[task_id] = task
        fire_immediately = bool(params.get("fire_immediately", False)) and not updating
        self._arm_scheduled_task(task_id, 0 if fire_immediately else interval_seconds)
        return json.dumps(
            {
                "id": task_id,
                "human_schedule": str(interval),
                "updated": updating,
            }
        )

    def execute_scheduler_delete(self, params: Dict[str, Any]) -> str:
        task_id = str(params.get("task_id", "")).strip()
        with self._scheduled_tasks_lock:
            task = self._scheduled_tasks.pop(task_id, None)
        if task and task.get("timer"):
            task["timer"].cancel()
        return json.dumps({"success": task is not None, "task_id": task_id})

    def execute_scheduler_list(self, _params: Dict[str, Any]) -> str:
        with self._scheduled_tasks_lock:
            tasks = [
                {
                    key: value
                    for key, value in task.items()
                    if key not in {"timer", "interval_seconds"}
                }
                for task in self._scheduled_tasks.values()
            ]
        return json.dumps({"tasks": tasks}, ensure_ascii=False)

    def execute_update_goal(self, params: Dict[str, Any]) -> str:
        completed = params.get("completed") is True
        blocked_reason = str(params.get("blocked_reason", "")).strip()
        message = str(params.get("message", "")).strip()
        if completed and blocked_reason:
            return "Error: completed and blocked_reason are mutually exclusive"
        status = "completed" if completed else "blocked" if blocked_reason else "active"
        self._goal_state = {
            "status": status,
            "message": message,
            "blocked_reason": blocked_reason,
            "updated_at": time.time(),
        }
        return json.dumps({"success": True, **self._goal_state}, ensure_ascii=False)

    def execute_read_pdf(self, params: Dict[str, Any]) -> str:
        """Read PDF or document file"""
        path = params.get("path", "")
        if not path:
            return "Error: path parameter required"

        try:
            expanded_path = self._resolve_project_path(path)

            # 检查文件类型
            if expanded_path.endswith(".docx") or expanded_path.endswith(".doc"):
                # 处理Word文档
                try:
                    from docx import Document

                    doc = Document(expanded_path)
                    text = ""
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text += para.text + "\n"

                    # 也提取表格内容
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join([cell.text for cell in row.cells])
                            text += row_text + "\n"

                    return f"Document contents:\n{text}"
                except ImportError:
                    return (
                        "Error: python-docx not installed. Try: pip install python-docx"
                    )

            elif expanded_path.endswith(".pdf"):
                # 处理PDF文件
                try:
                    import PyPDF2

                    with open(expanded_path, "rb") as file:
                        reader = PyPDF2.PdfReader(file)
                        text = ""
                        for page in reader.pages:  # Read all pages
                            text += page.extract_text() + "\n"
                    return f"PDF contents:\n{text}"
                except ImportError:
                    return "Error: PyPDF2 not installed. Try: pip install PyPDF2"

            else:
                return f"Error: Unsupported file format. Supported: .pdf, .docx, .doc"

        except Exception as e:
            return f"Error reading document: {str(e)}"

    def execute_read_url(self, params: Dict[str, Any]) -> str:
        """Read and extract content from a URL (bounded fetch)."""
        url = params.get("url", "")
        if not url:
            return "Error: url parameter required"

        outcome = fetch_url(url)
        if not outcome.ok:
            return f"Error: {outcome.error}"

        content = outcome.text

        if outcome.kind == "html":
            # Try using BeautifulSoup if available
            try:
                from bs4 import BeautifulSoup

                soup = BeautifulSoup(content, "html.parser")

                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()

                # Get text
                text = soup.get_text()

                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (
                    phrase.strip() for line in lines for phrase in line.split("  ")
                )
                text = "\n".join(chunk for chunk in chunks if chunk)

                if text:
                    content = text
                else:
                    return "Error: 无法从 URL 提取内容"

            except ImportError:
                # If BeautifulSoup not available, try simple regex extraction
                import re

                # Remove script and style tags
                content = re.sub(
                    r"<script[^>]*>.*?</script>", "", content, flags=re.DOTALL
                )
                content = re.sub(
                    r"<style[^>]*>.*?</style>", "", content, flags=re.DOTALL
                )

                # Remove HTML tags
                content = re.sub(r"<[^>]+>", "", content)

                # Clean up whitespace
                content = re.sub(r"\s+", " ", content).strip()

                if not content:
                    return "Error: 无法从 URL 提取内容"

        if outcome.truncated:
            content += "\n\n(内容已截断，请访问更具体的 URL 或章节获取全文)"

        return f"URL 内容:\n{content}"

    def execute_send_file(self, params: Dict[str, Any]) -> str:
        """Send a file to the user via Feishu"""
        import asyncio
        from agent.bus.events import OutboundMessage

        file_path = params.get("path", "")

        if not file_path:
            return "Error: 必须指定文件路径"
        file_path = self._resolve_project_path(file_path)

        import os

        if not os.path.isfile(file_path):
            return f"Error: 文件不存在 - {file_path}"

        # Get the current executor context to access bus and chat info
        # This is a bit hacky but necessary for the current architecture
        import inspect

        frame = inspect.currentframe()
        executor = None

        # Walk up the stack to find NaturalTaskExecutor
        while frame:
            if "self" in frame.f_locals:
                obj = frame.f_locals["self"]
                if hasattr(obj, "bus") and hasattr(obj, "current_chat_id"):
                    executor = obj
                    break
            frame = frame.f_back

        if not executor or not executor.bus or not executor.current_chat_id:
            return "Error: 无法发送文件 - 未在网关模式下运行"

        try:
            # Create outbound message with file path
            msg = OutboundMessage(
                channel=executor.current_channel or "feishu",
                chat_id=executor.current_chat_id,
                content=file_path,  # Pass file path as content
            )

            # Send via bus (non-blocking)
            asyncio.create_task(executor.bus.publish_outbound(msg))

            file_name = os.path.basename(file_path)
            file_size = os.path.getsize(file_path)
            return f"✅ 文件已发送: {file_name} ({file_size} bytes)"

        except Exception as e:
            return f"Error: 发送文件失败 - {str(e)}"

    def execute_generate_pdf(self, params: Dict[str, Any]) -> str:
        """Generate PDF from various formats"""
        input_path = params.get("input_path", "")
        output_path = params.get("output_path", "")
        format_type = params.get("format", "")

        # Parameter validation
        if not input_path or not output_path:
            return "Error: input_path and output_path parameters required"
        protection_error = self._mutation_path_error(output_path)
        if protection_error:
            return protection_error

        # Auto-detect format from file extension if not specified
        if not format_type:
            if input_path.lower().endswith(".md"):
                format_type = "markdown"
            elif input_path.lower().endswith(".html"):
                format_type = "html"
            elif input_path.lower().endswith(".docx") or input_path.lower().endswith(
                ".doc"
            ):
                format_type = "docx"
            else:
                format_type = "text"

        if format_type not in ["markdown", "text", "html", "docx"]:
            return f"Error: Unsupported format '{format_type}'. Supported: markdown, text, html, docx"

        # Path expansion
        expanded_input = self._resolve_project_path(input_path)
        expanded_output = self._resolve_project_path(output_path)

        # Call PDF tool
        success, message = self.pdf_tool.generate_pdf(
            expanded_input, expanded_output, format_type
        )

        if success:
            return f"✅ {message}"
        else:
            return message

    def execute_load_skill(
        self,
        params: Dict[str, Any],
        conversation_id: Optional[str] = None,
        message_id: Optional[str] = None,
    ) -> str:
        """Load a skill's complete content"""
        if not self.skill_tool:
            return "Error: Skill tool not initialized"

        skill_name = params.get("skill_name", "").strip()
        if not skill_name:
            return "Error: skill_name parameter required"

        success, content = self.skill_tool.load_skill(skill_name)
        if success:
            key = self._plan_key(conversation_id, message_id)
            with self._loaded_skills_lock:
                self._loaded_skills.setdefault(key, set()).add(skill_name)
        return content

    def execute_question(self, params: Dict[str, Any]) -> str:
        """Execute question - ask user for input"""
        questions = params.get("questions", [])
        if not questions:
            return "Error: questions parameter required"

        # Format question for user
        formatted = []
        for q in questions:
            header = q.get("header", "Question")
            question_text = q.get("question", "")
            options = q.get("options", [])

            formatted.append(f"【{header}】")
            formatted.append(question_text)
            if options:
                formatted.append("选项:")
                for opt in options:
                    formatted.append(
                        f"  - {opt.get('label', '')}: {opt.get('description', '')}"
                    )
            formatted.append("")

        return "请回复你的选择:\n" + "\n".join(formatted)

    def execute_project_preview(
        self,
        params: Dict[str, Any],
        conversation_id: Optional[str] = None,
        message_id: Optional[str] = None,
    ) -> str:
        """Execute a persistent project preview lifecycle action."""
        action = str(params.get("action", "")).strip().lower()
        preview_id = str(params.get("preview_id", "")).strip()

        if action == "start":
            result = self.preview_manager.start(
                command=params.get("command", ""),
                workdir=params.get("workdir", "."),
                name=params.get("name"),
                port=params.get("port", 0),
                health_path=params.get("health_path", "/"),
                entry_path=params.get("entry_path"),
                startup_timeout=params.get("startup_timeout", 20),
                conversation_id=conversation_id,
                message_id=message_id,
            )
        elif action == "status":
            result = self.preview_manager.status(
                preview_id=preview_id or None,
                conversation_id=conversation_id,
            )
        elif action == "stop":
            if not preview_id:
                return "Error: preview_id parameter required for stop"
            result = self.preview_manager.stop(preview_id)
        else:
            return "Error: action must be start, status, or stop"

        return json.dumps(result, ensure_ascii=False)
